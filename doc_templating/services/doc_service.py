import logging
import re
from dataclasses import dataclass
from io import BytesIO
from typing import Mapping

from docx import Document
from docx.document import Document as Document_cls
from docx.opc.exceptions import PackageNotFoundError

from doc_templating import config

log = logging.getLogger(__name__)


@dataclass
class Field:
    full_text: str
    key: str
    default_value: str
    typ: str


def load_template():
    try:
        template = Document(config["szablon_plik"].get(str))
    except PackageNotFoundError as err:
        raise FileNotFoundError(str(err)) from err
    return template


def get_keys(doc: Document_cls = None) -> dict[str, Field]:
    """Get keys and their default values.

    For a key specified like ``{my key}``, will assume default value of an empty string.
    For a key specified like ``{my key = my value}``, the key will be "my key"
    and the default will be "my value" (both without surrounding whitespace).

    Returns
    -------
    res
        Dict key -> Field
    """
    if doc is None:
        doc = load_template()
    res = {}
    for p in doc.paragraphs:
        for r in p.runs:
            for match in re.finditer(r"\{.*?}", r.text):
                full_text = match.group().strip("{}")
                if "=" in full_text:
                    key, default = map(str.strip, full_text.split("="))
                else:
                    key, default = full_text.strip(), ""
                if ":" in key:
                    typ, key = map(str.strip, key.split(":"))
                else:
                    typ, key = "text", key.strip()
                res[key] = Field(full_text, key, default, typ)
    if "cena całkowita" in res and "stawka za wizytę" in res and "liczba wizyt" in res:
        res["cena całkowita"].default_value = "---auto---"
    return res


def substitute(substitutions: Mapping[str, str], doc: Document_cls) -> Document_cls:
    for p in doc.paragraphs:
        for run in p.runs:
            for key, value in substitutions.items():
                run.text = run.text.replace(f"{{{key}}}", value)
    return doc


def substitute_and_stream(substitutions: Mapping[str, str], doc: Document_cls = None):
    if doc is None:
        doc = load_template()
    out = substitute(substitutions, doc)
    target_stream = BytesIO()
    out.save(target_stream)
    target_stream.seek(0)
    return target_stream


def preprocess_substitutions(
    substitutions: Mapping[str, str], fields: Mapping[str, Field] = None
):
    if fields is None:
        fields = get_keys(load_template())
    res = dict(substitutions)
    if (
        fields["cena całkowita"].full_text in res
        and fields["stawka za wizytę"].full_text in res
        and fields["liczba wizyt"].full_text in res
        and res[fields["cena całkowita"].full_text] == "---auto---"
    ):
        try:
            res[fields["cena całkowita"].full_text] = str(
                cena := (
                    int(res[fields["stawka za wizytę"].full_text])
                    * int(res[fields["liczba wizyt"].full_text])
                )
            )
        except ValueError as err:
            raise ValueError("Upewnij się, że w polch liczbowych wpisane są liczby.")
        log.debug(f"Obliczono cenę całkowitą: {cena} zł")
    else:
        log.debug("Brak danych do obliczenia ceny.")
    return res


def fill_missing_substitutions(substitutions: Mapping[str, str]) -> dict[str, str]:
    all_keys = get_keys()
    res = dict(substitutions)
    for key, field in all_keys.items():
        if field.full_text not in res and field.typ == "checkbox":
            res[field.full_text] = ""
    return res
